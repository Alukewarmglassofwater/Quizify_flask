from app import app
from flask import Flask, render_template, redirect, url_for, request, session, flash, g, send_from_directory, jsonify, render_template_string
from functools import wraps
from docx import Document
import sqlite3
import re

# Connect to .db file
DATABASE = 'app/instance/database.db'
# list of answers selected by the user
user_answer_list = []
correct_answer_list = []

@app.route('/favicon.ico')
def favicon():
    return send_from_directory('static/favicon', 'favicon.ico', mimetype='image/x-icon')

def login_required(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'logged_in' in session:
            return f(*args, **kwargs)
        else:
            return redirect(url_for('login'))
    return wrap



@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    mesage = ''
    if request.method == 'POST' and 'email' in request.form and 'password' in request.form:
        email = request.form['email']
        password = request.form['password']        
        conn = sqlite3.connect(DATABASE)
        conn.row_factory = sqlite3.Row  
        c = conn.cursor()
        # Select all user to check correct email and pass
        c.execute('SELECT * FROM user WHERE email = ? AND password = ?', (email, password, ))
        user = c.fetchone()
        if user:
            session['logged_in'] = True
            session['ID'] = user['ID']
            session['name'] = user['name']
            session['email'] = user['email']
            mesage = 'Logged in successfully !'
            return redirect(url_for('home'))
        else:
            mesage = 'Please enter correct email / password !'
    return render_template('login.html', mesage = mesage)


@app.route('/logout')
@login_required
def logout():
    session.pop('logged_in', None)
    return redirect('/login')



@app.route('/home')
@login_required
def home():
    # get the current user email
    email = session['email']
    conn = sqlite3.connect(DATABASE)
    c = conn.cursor()

    # Fetch the leaderboard
    c.execute('SELECT * FROM userquizscore ORDER BY score DESC')
    leaderboard = c.fetchall()

    # Fetch the user profile if no score then score = 0
    c.execute('''
        SELECT user.ID, user.name, user.email, IFNULL(userquizscore.score, 0) as score
        FROM user
        LEFT JOIN userquizscore
        ON user.email = userquizscore.email
        WHERE user.email = ?
    ''', (email,))
    userProfile = c.fetchone()
    conn.commit()
    conn.close()

    return render_template('home.html', leaderboard=leaderboard, userProfile=userProfile)


@app.route('/register', methods = ['GET', 'POST'])
def register():
    mesage = ''
    # Check if user fill the form or not or is it the correct format
    if request.method == 'POST' and 'name' in request.form and 'password' in request.form and 'email' in request.form :
        userName = request.form['name']
        password = request.form['password']
        email = request.form['email']
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute('SELECT * FROM user WHERE email = ?', (email,))
        account = c.fetchone()
        if account:
            mesage = 'Account already exists !'
        elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
            mesage = 'Invalid email address !'
        elif not userName or not password or not email:
            mesage = 'Please fill out the form !'
        else:
            c.execute('''INSERT INTO user (name, email, password)VALUES (?, ?, ?)''', 
                      (userName, email, password))
            conn.commit()
            conn.close() 
            mesage = 'You have successfully registered !'
    elif request.method == 'POST':
        mesage = 'Please fill out the form !'
    return render_template('register.html', mesage = mesage)

#possibly mcq quiz logic?? not really sure how it works??
# Fetch questions and answers from the database
def get_question(index):
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM multichoice LIMIT 1 OFFSET ?", (index,))
    question_data = cursor.fetchone()
    conn.close()
    return question_data

from flask import jsonify

@app.route('/quizmcq', methods=['GET', 'POST'])
def quiz():
    if request.method == 'POST':
        # Check if any radio button is selected
        if 'answer' not in request.form:
            # If no radio button is selected, return an alert box using JavaScript
            return render_template_string('<script>alert("Please select an answer."); window.history.back();</script>')

        # Get selected answer
        selected_answer = request.form['answer']
        user_answer_list.append(selected_answer)
        print("list: ", user_answer_list)

        # Increment the question index for the next question
        index = int(request.args.get('index', 0)) + 1
        return redirect(url_for('quiz', index=index))

    # Get the current question index from query parameter
    index = int(request.args.get('index', 0))

    # Fetch the question and answer options from the database
    question_data = get_question(index)
    
    if question_data:
        question, answer1, answer2, answer3, answer4, correct_answer = question_data
        options = [answer1, answer2, answer3, answer4]
        correct_answer_list.append(correct_answer)
        return render_template('quizmcq.html', question=question, options=options, index=index, correct_answer=correct_answer)
    else:
        # No more questions, quiz completed
        score_sum = 0
        for index, item in enumerate(user_answer_list):
            if user_answer_list[index] == correct_answer_list[index]:
                score_sum = score_sum + 1
        
        # Store quiz score into the database
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()

        # Check if the email already exists in the userquizscore table
        email = session.get("email")
        c.execute('''SELECT * FROM userquizscore WHERE email = ?''', (email,))
        existing_record = c.fetchone()

        if existing_record:
            # Email already exists, update the existing record
            c.execute('''UPDATE userquizscore SET name = ?, score = ? WHERE email = ?''', 
                        (session.get("name"), score_sum, email))
        else:
            # Insert a new record
            c.execute('''INSERT INTO userquizscore (name, email, score) VALUES (?, ?, ?)''', 
                        (session.get("name"), email, score_sum))

        conn.commit()
        conn.close()
        
        user_answer_list.clear()
        correct_answer_list.clear()

        return render_template('quizcompleted.html', score_sum=score_sum)

doc = Document()
def export_to_word(question, answer):
    
    doc.add_paragraph('Question: ' + question)
    doc.add_paragraph('Answer: ' + answer)


@app.route('/quizsa', methods=['GET', 'POST'])
def quizsa():
    global doc
    if request.method == 'POST':
        # Get the submitted question index
        question_index = int(request.form['index'])
        
        # Check for button actions
        if 'submit' in request.form:
            print("submit clicked")
            question_index += 1
            
        if 'action' in request.form:
            action = request.form['action']
            if action == 'forward':
                # Move to the next question
                question_index += 1
            elif action == 'backward':
                # Move to the previous question
                if question_index > 0:
                    question_index -= 1
        
        # Connect to the database
        conn = sqlite3.connect(DATABASE)
        cursor = conn.cursor()

        # Fetch short-answer questions from the database
        cursor.execute("SELECT question FROM shortanswerquestions")
        questions = [row[0] for row in cursor.fetchall()]
        # Check if the index is valid
        if 0 <= question_index < len(questions):
            # Fetch the current question
            current_question = questions[question_index]
            answer = request.form['answer']
            print("Question: " + current_question + " Answer: " + answer + " Current Index: " + str(question_index))            
            export_to_word(current_question, answer)
            # Close the database connection
            conn.close()

            return render_template('quizsa.html', title='Shortanswer', question=current_question, index=question_index)
        # Save the document
        doc.save('quiz_answers.docx')  
        doc = Document()
        # No more questions, quiz completed
        conn.close()
        return render_template('quizcompletedSA.html')
    else:
        # Connect to the database
        conn = sqlite3.connect(DATABASE)
        cursor = conn.cursor()

        # Fetch the first short-answer question from the database
        cursor.execute("SELECT question FROM shortanswerquestions LIMIT 1")
        first_question = cursor.fetchone()[0]

        # Close the database connection
        conn.close()

        return render_template('quizsa.html', title='Shortanswer', question=first_question, index=0)



@app.route('/quizcompleted')
def quiz_completed():
    return render_template('quizcompleted.html')

@app.route('/quizcompletedSA')
def quiz_completed_SA():
    return render_template('quizcompletedSA.html')

@app.route('/retakingquiz')
def retakingquiz():
    return render_template('retake.html')

from flask import render_template
